"""
RapidRAG — ingestion.py
Ingestion orchestrator: scrape/PDF → chunk → embed → store vectors.
Includes re-ingestion locking and embedding version tracking.
"""

import os
import asyncio
import logging
from pathlib import Path

import hashlib
import fitz  # PyMuPDF
import docx  # python-docx
import pandas as pd  # pandas

from config import settings, IngestionResult, Chunk
from scraper import WebScraper
from chunker import Chunker
from embedder import Embedder
from client_manager import ClientManager
from errors import IngestionLockError, EmbeddingMismatchError

logger = logging.getLogger(__name__)


class IngestionPipeline:
    """
    Orchestrates the full ingestion lifecycle:
      Website URL / PDF → scrape/extract → chunk → embed → store vectors

    Safety features:
    - Per-client asyncio.Lock prevents concurrent ingestion corruption
    - Embedding version tracking detects model mismatches
    - Stores raw + clean text in client documents/ folder
    """

    def __init__(
        self,
        scraper: WebScraper,
        chunker: Chunker,
        embedder: Embedder,
        client_manager: ClientManager,
    ):
        self.scraper = scraper
        self.chunker = chunker
        self.embedder = embedder
        self.client_manager = client_manager

        # Per-client ingestion locks (Missing Fix 3)
        self._locks: dict[str, asyncio.Lock] = {}

    # ─────────────────────────────────────────
    # Lock Management
    # ─────────────────────────────────────────

    def _get_lock(self, client_id: str) -> asyncio.Lock:
        """Get or create per-client ingestion lock."""
        if client_id not in self._locks:
            self._locks[client_id] = asyncio.Lock()
        return self._locks[client_id]

    def is_ingesting(self, client_id: str) -> bool:
        """Check if ingestion is currently running for a client."""
        if client_id not in self._locks:
            return False
        return self._locks[client_id].locked()

    # ─────────────────────────────────────────
    # Website Ingestion
    # ─────────────────────────────────────────

    async def ingest_website(
        self, client_id: str, url: str
    ) -> IngestionResult:
        """
        Full website ingestion pipeline:
        1. Acquire lock (prevent concurrent ingestion)
        2. Check embedding version compatibility
        3. Scrape website
        4. Save raw text to documents/
        5. Chunk all scraped text
        6. Batch embed chunks
        7. Store vectors in ChromaDB
        8. Write embedding version stamp
        """
        lock = self._get_lock(client_id)

        if lock.locked():
            raise IngestionLockError(
                f"Ingestion already in progress for '{client_id}'. "
                f"Please wait for it to complete."
            )

        async with lock:
            errors = []
            docs_dir = os.path.join(
                settings.base_client_path, client_id, "documents"
            )
            os.makedirs(docs_dir, exist_ok=True)

            # Step 1: Scrape website
            logger.info(f"[{client_id}] Scraping website: {url}")
            try:
                pages = await self.scraper.scrape_site(url)
            except Exception as e:
                logger.error(f"[{client_id}] Scraping failed: {e}")
                return IngestionResult(
                    status="failed",
                    errors=[f"Scraping error: {str(e)}"],
                )

            if not pages:
                return IngestionResult(
                    status="completed",
                    pages_scraped=0,
                    errors=["No pages scraped — check URL"],
                )

            # Step 2: Save scraped text to documents/
            documents = []
            for page in pages:
                if not page.clean_text.strip():
                    continue

                # Save clean text
                slug = self._url_to_slug(page.url)
                text_path = os.path.join(docs_dir, f"{slug}.txt")
                try:
                    with open(text_path, "w", encoding="utf-8") as f:
                        f.write(page.clean_text)
                except Exception as e:
                    errors.append(f"Failed to save {slug}: {e}")
                    continue

                documents.append({
                    "text": page.clean_text,
                    "source": slug,
                })

            logger.info(
                f"[{client_id}] Saved {len(documents)} pages to documents/"
            )

            # Step 3-5: Chunk → Embed → Store
            result = await self._process_documents(
                client_id, documents, len(pages)
            )
            result.pages_scraped = len(pages)
            result.errors.extend(errors)

            return result

    # ─────────────────────────────────────────
    # File Ingestion (Generic)
    # ─────────────────────────────────────────

    async def ingest_file(
        self, client_id: str, file_path: str, filename: str = ""
    ) -> IngestionResult:
        """
        Generic File Ingestion Pipeline for PDF, DOCX, CSV, XLSX:
        1. Acquire lock
        2. Detect extension and extract text natively
        3. Save extracted text
        4. Chunk → embed → store
        """
        lock = self._get_lock(client_id)

        if lock.locked():
            raise IngestionLockError(
                f"Ingestion already in progress for '{client_id}'"
            )

        async with lock:
            docs_dir = os.path.join(
                settings.base_client_path, client_id, "documents"
            )
            os.makedirs(docs_dir, exist_ok=True)

            # Step 1: Detect and Extract Text
            source_name = filename or os.path.basename(file_path)
            ext = Path(source_name).suffix.lower()
            logger.info(f"[{client_id}] Extracting file [{ext}]: {source_name}")

            try:
                if ext == ".pdf":
                    text = await asyncio.to_thread(self._extract_pdf_text, file_path)
                elif ext == ".docx":
                    text = await asyncio.to_thread(self._extract_docx_text, file_path)
                elif ext in [".xls", ".xlsx", ".csv"]:
                    text = await asyncio.to_thread(self._extract_spreadsheet_text, file_path, ext)
                else:
                    return IngestionResult(
                        status="failed",
                        errors=[f"Unsupported file type: {ext}"]
                    )
            except Exception as e:
                logger.error(f"[{client_id}] File extraction failed: {e}")
                return IngestionResult(
                    status="failed",
                    errors=[f"File extraction error: {str(e)}"],
                )

            if not text.strip():
                return IngestionResult(
                    status="completed",
                    pages_parsed=0,
                    errors=[f"File appears empty or unreadable: {source_name}"],
                )

            # Step 2: Save extracted text
            slug = Path(source_name).stem
            slug = "".join(c if c.isalnum() or c in "-_" else "_" for c in slug)
            text_path = os.path.join(docs_dir, f"{slug}.txt")
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(text)

            documents = [{"text": text, "source": slug}]

            # Step 3-5: Chunk → Embed → Store
            result = await self._process_documents(
                client_id, documents, 0
            )
            result.pages_parsed = 1
            return result

    # ─────────────────────────────────────────
    # Text Ingestion (Direct Text Append)
    # ─────────────────────────────────────────

    async def ingest_text(
        self, client_id: str, text: str, source: str = "manual-text"
    ) -> IngestionResult:
        """
        Ingest raw text and append to existing knowledge base.
        Safely acquires locks and writes sequentially.
        """
        lock = self._get_lock(client_id)

        if lock.locked():
            raise IngestionLockError(
                f"Ingestion already in progress for '{client_id}'"
            )

        async with lock:
            docs_dir = os.path.join(
                settings.base_client_path, client_id, "documents"
            )
            os.makedirs(docs_dir, exist_ok=True)

            safe_source = "".join(c if c.isalnum() or c in "-_" else "_" for c in source)
            text_path = os.path.join(docs_dir, f"{safe_source}.txt")
            
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(text)
            
            documents = [{"text": text, "source": safe_source}]
            
            # Use append=True so we don't wipe the DB
            result = await self._process_documents(
                client_id, documents, 0, append=True
            )
            result.pages_parsed = 1
            return result

    # ─────────────────────────────────────────
    # Shared Processing Pipeline
    # ─────────────────────────────────────────

    async def _process_documents(
        self,
        client_id: str,
        documents: list[dict[str, str]],
        pages_scraped: int,
        append: bool = False,
    ) -> IngestionResult:
        """
        Shared pipeline: chunk → embed → store.
        Used by both website and PDF ingestion.
        """
        # Step 3: Chunk all documents
        logger.info(f"[{client_id}] Chunking {len(documents)} documents")
        chunks = self.chunker.chunk_documents(documents)

        if not chunks:
            return IngestionResult(
                status="completed",
                pages_scraped=pages_scraped,
                chunks_created=0,
                errors=["No chunks created — text may be too short"],
            )

        # Step 4: Batch embed all chunks
        logger.info(f"[{client_id}] Embedding {len(chunks)} chunks")
        chunk_texts = [c.text for c in chunks]
        embeddings = await self.embedder.embed_batch(chunk_texts)

        # Step 5: Store vectors in ChromaDB
        logger.info(f"[{client_id}] Storing vectors in ChromaDB")
        await self._store_vectors(client_id, chunks, embeddings, append=append)

        # Step 6: Write embedding version stamp
        self._write_embedding_version(client_id)

        result = IngestionResult(
            status="completed",
            pages_scraped=pages_scraped,
            chunks_created=len(chunks),
            vectors_stored=len(embeddings),
            embedding_model=self.embedder.model_name,
        )

        logger.info(
            f"[{client_id}] Ingestion complete: "
            f"{result.chunks_created} chunks, "
            f"{result.vectors_stored} vectors"
        )
        return result

    # ─────────────────────────────────────────
    # Vector Storage
    # ─────────────────────────────────────────

    async def _store_vectors(
        self,
        client_id: str,
        chunks: list[Chunk],
        embeddings: list[list[float]],
        append: bool = False,
    ) -> None:
        """Store chunks + embeddings in the client's ChromaDB collection."""
        collection = self.client_manager.get_vector_store(client_id)

        # Clear existing vectors (full re-ingestion replaces all) unless appending
        existing_count = await asyncio.to_thread(collection.count)
        if not append and existing_count > 0:
            # Fetch IDs to delete them (using empty get() to retrieve only IDs in recent ChromaDB)
            existing = await asyncio.to_thread(collection.get)
            if existing["ids"]:
                # Delete in batches to avoid memory spikes
                batch_del = 5000
                for i in range(0, len(existing["ids"]), batch_del):
                    await asyncio.to_thread(collection.delete, ids=existing["ids"][i:i + batch_del])
            logger.info(
                f"[{client_id}] Cleared {existing_count} existing vectors"
            )

        # Prepare batch data (using content hashing + chunk index for IDs to be safe on append)
        unique_ids = []
        unique_docs = []
        unique_embs = []
        unique_metas = []
        seen_ids = set()
        
        for c, emb in zip(chunks, embeddings):
            chunk_id = f"{client_id}_{hashlib.sha256(f'{c.source}_{c.chunk_index}_{c.text}'.encode()).hexdigest()[:16]}"
            if chunk_id not in seen_ids:
                seen_ids.add(chunk_id)
                unique_ids.append(chunk_id)
                unique_docs.append(c.text)
                unique_embs.append(emb)
                unique_metas.append({
                    "source": c.source,
                    "chunk_index": c.chunk_index,
                })

        # Upsert in batches (ChromaDB handles large batches fine, but be safe)
        batch_size = 500
        for i in range(0, len(unique_ids), batch_size):
            end = i + batch_size
            await asyncio.to_thread(
                collection.upsert,
                ids=unique_ids[i:end],
                documents=unique_docs[i:end],
                embeddings=unique_embs[i:end],
                metadatas=unique_metas[i:end],
            )

        logger.info(
            f"[{client_id}] Stored {len(unique_ids)} vectors in ChromaDB (deduped from {len(chunks)})"
        )

    # ─────────────────────────────────────────
    # Embedding Versioning (Missing Fix 1)
    # ─────────────────────────────────────────

    def _write_embedding_version(self, client_id: str) -> None:
        """Write current embedding model name after successful ingestion."""
        version_path = os.path.join(
            settings.base_client_path, client_id,
            "vectors", "embedding_model.txt",
        )
        with open(version_path, "w", encoding="utf-8") as f:
            f.write(self.embedder.model_name)

    def check_embedding_version(self, client_id: str) -> bool:
        """
        Check if stored vectors match current embedding model.
        Returns True if compatible, False if re-ingestion needed.
        """
        version_path = os.path.join(
            settings.base_client_path, client_id,
            "vectors", "embedding_model.txt",
        )
        if not os.path.exists(version_path):
            return False  # no version file = legacy, needs re-ingest

        with open(version_path, "r", encoding="utf-8") as f:
            stored_model = f.read().strip()

        return stored_model == self.embedder.model_name

    # ─────────────────────────────────────────
    # File Text Extraction Helpers
    # ─────────────────────────────────────────

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from a PDF file using PyMuPDF (fitz)."""
        doc = fitz.open(file_path)
        total_pages = len(doc)  # Save before close
        text_parts = []

        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

        doc.close()

        full_text = "\n\n".join(text_parts)
        logger.info(
            f"PDF extracted: {total_pages} pages, "
            f"{len(full_text)} chars"
        )
        return full_text

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from a Word Document (.docx) using python-docx."""
        document = docx.Document(file_path)
        text_parts = []
        
        for para in document.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
                
        # Also extract basic table text
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"DOCX extracted: {len(full_text)} chars")
        return full_text

    def _extract_spreadsheet_text(self, file_path: str, ext: str) -> str:
        """Extract text from Excel (.xlsx, .xls) or CSV files using pandas."""
        text_parts = []
        
        if ext == '.csv':
            # Parse CSV
            df = pd.read_csv(file_path)
            text_parts.append(df.to_string(index=False))
        else:
            # Parse Excel (all sheets)
            xl = pd.ExcelFile(file_path)
            for sheet_name in xl.sheet_names:
                df = xl.parse(sheet_name)
                # Drop entirely empty rows/columns
                df = df.dropna(how='all').dropna(axis=1, how='all')
                if not df.empty:
                    text_parts.append(f"--- Sheet: {sheet_name} ---")
                    text_parts.append(df.to_string(index=False))

        full_text = "\n\n".join(text_parts)
        logger.info(f"Spreadsheet extracted: {len(full_text)} chars")
        return full_text

    # ─────────────────────────────────────────
    # Utilities
    # ─────────────────────────────────────────

    @staticmethod
    def _url_to_slug(url: str) -> str:
        """Convert a URL to a filesystem-safe slug (includes domain to avoid collisions)."""
        from urllib.parse import urlparse
        parsed = urlparse(url)
        domain = parsed.netloc.replace("www.", "").split(":")[0]  # strip www and port
        path = parsed.path.strip("/").replace("/", "_") or "homepage"
        raw = f"{domain}_{path}" if path != "homepage" else f"{domain}_homepage"
        # Remove special characters
        slug = "".join(c if c.isalnum() or c == "_" else "_" for c in raw)
        # Collapse multiple underscores
        while "__" in slug:
            slug = slug.replace("__", "_")
        return slug[:120]  # cap length
